from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Pipeline_Risk_Project_Guide.docx"

def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    properties.append(element)

def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)

def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)

def add_heading(document, text, level=1):
    document.add_heading(text, level=level)

def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        shade(table.rows[0].cells[index], "173A4A")
        for run in table.rows[0].cells[index].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    document.add_paragraph()
    return table

document = Document()
section = document.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

styles = document.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].paragraph_format.space_after = Pt(6)
for name, size, color in [("Title", 26, "173A4A"), ("Heading 1", 17, "176B87"), ("Heading 2", 13, "173A4A")]:
    styles[name].font.name = "Aptos Display"
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(color)

header = section.header.paragraphs[0]
header.text = "PIPELINE RISK MONITOR  |  TEAM GUIDE"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header.runs[0].font.size = Pt(8)
header.runs[0].font.color.rgb = RGBColor(100, 110, 120)

footer = section.footer.paragraphs[0]
footer.text = "SCADA Pipeline Risk Assessment Project"
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor(100, 110, 120)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Pipeline Risk Monitor")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(23, 107, 135)
subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A simple guide for teammates")
run.italic = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 90, 100)

intro = document.add_paragraph()
intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
intro.add_run("This project analyzes SCADA readings from a gas pipeline network and estimates leak or abnormal-operation risk for each pipeline segment.").bold = True
document.add_paragraph("The project currently runs on a synthetic/sample SCADA dataset. It has been prepared for future Titas Gas integration, but it is not connected to live Titas data.")

add_heading(document, "1. What Is This Project?", 1)
document.add_paragraph("The Pipeline Risk Monitor is a machine-learning dashboard for monitoring pressure, flow, temperature, equipment states, energy use, and alarm signals. It turns sensor readings into a 0–100% risk score and an alert level.")
add_bullets(document, [
    "Normal: the reading is within the model's normal-risk range.",
    "Warning: the reading may need operator attention.",
    "Critical Leak Threat: the reading has a high predicted risk and should be investigated.",
])
document.add_paragraph("The model is an early-warning tool. It does not replace an operator, a safety system, a pressure test, a field inspection, or an official leak investigation.")

add_heading(document, "2. What Is Included?", 1)
add_table(document, ["Area", "What it contains"], [
    ("Machine learning", "Isolation Forest anomaly detection plus XGBoost or Random Forest classification."),
    ("Feature engineering", "Rolling pressure, flow, temperature, pump-speed, and energy statistics; gradients; anomaly indexes; state-change flags; and IQR sensor-outlier flags."),
    ("Backend", "Flask application with model APIs, data validation, CSV export, topology data, and alert actions."),
    ("Frontend", "SCADA-style browser dashboard built with HTML, CSS, JavaScript, Chart.js, and SVG."),
    ("Operations", "Risk timeline, telemetry charts, segment overview, detection log, scenario presets, and alert acknowledgement."),
    ("Data preparation", "Titas-oriented configuration, unit conversion, timestamp normalization, range validation, and quality reporting."),
])

add_heading(document, "3. Main Technologies", 1)
add_bullets(document, [
    "Python: application and machine-learning code.",
    "Flask: web server and REST API.",
    "Pandas and NumPy: data preparation and numerical operations.",
    "Scikit-learn: time-series evaluation, Isolation Forest, scaling, and fallback Random Forest model.",
    "XGBoost: supervised risk classifier when available.",
    "SHAP: explanation of the most influential features for a prediction.",
    "Chart.js and browser JavaScript: charts and dashboard interactions.",
    "SQLite: persistent alert acknowledgement and resolution state.",
    "YAML: SCADA tag mapping, unit rules, and configuration.",
])

add_heading(document, "4. Project Folder Structure", 1)
add_table(document, ["Folder/file", "Purpose"], [
    ("data/raw/", "Original source SCADA sample, preserved for provenance."),
    ("data/processed/", "Collision-repaired sample used by the local demo model."),
    ("data/runtime/", "Local runtime state such as SQLite alert records; not committed to Git."),
    ("config/titas_scada.yml", "Canonical fields, source tag mapping, units, ranges, and required data."),
    ("config/demo_topology.yml", "Demo schematic layout. It is not a geographic Titas map."),
    ("webapp/app.py", "Flask entry point, startup loading, API routes, and dashboard serving."),
    ("webapp/model_pipeline.py", "Reusable risk model, feature engineering, scoring, simulation, and artifact loading."),
    ("webapp/scada_adapter.py", "Input mapping, unit conversion, timestamp normalization, and quality checks."),
    ("webapp/alert_store.py", "SQLite storage for acknowledge, resolve, and false-positive alert states."),
    ("webapp/templates/", "Dashboard HTML template."),
    ("webapp/static/", "JavaScript, CSS, and generated dashboard plots."),
    ("scripts/train_model.py", "Validated model training and artifact creation."),
    ("scripts/audit_data.py", "Data audit before training or live shadow mode."),
    ("tests/", "Adapter and alert persistence tests."),
    ("docs/", "Research and integration documentation."),
])

add_heading(document, "5. Backend Architecture", 1)
document.add_paragraph("The backend follows a simple pipeline:")
add_numbered(document, [
    "Load configuration and the SCADA dataset.",
    "Pass the data through the SCADA adapter.",
    "Normalize timestamps and units, validate ranges, and report quality issues.",
    "Load a reviewed model artifact when PIPELINE_RISK_MODEL_PATH is set.",
    "Otherwise, train a development model at startup.",
    "Expose prediction, simulation, operations, topology, export, and alert APIs.",
    "Serve the browser dashboard from Flask.",
])
document.add_paragraph("In a production deployment, training should happen separately. The dashboard should load a reviewed model artifact instead of retraining whenever the server restarts.")

add_heading(document, "6. Important API Routes", 1)
add_table(document, ["Route", "Purpose"], [
    ("GET /", "Open the dashboard."),
    ("GET /api/meta", "Dataset, model, and data-quality metrics."),
    ("GET /api/operations", "Current operations summary, detections, and segment data."),
    ("GET /api/topology", "Demo segment topology and risk colors."),
    ("GET /api/simulate/<segment_id>", "Replay a segment's historical readings through the model."),
    ("POST /api/predict", "Score a manual what-if reading."),
    ("GET /api/export/data", "Download all SCADA data as CSV."),
    ("POST /api/alerts/<alert_key>/acknowledge", "Acknowledge an alert and store it in SQLite."),
    ("POST /api/alerts/<alert_key>/resolve", "Resolve an acknowledged alert."),
])

add_heading(document, "7. How the Model Works", 1)
document.add_paragraph("The model combines two signals:")
add_bullets(document, [
    "Anomaly score: Isolation Forest estimates how unusual the operating conditions are compared with training data.",
    "Classifier probability: XGBoost or Random Forest estimates the probability of an abnormal or incident condition using labeled history.",
])
document.add_paragraph("The final risk score blends these signals:")
formula = document.add_paragraph()
formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
formula.add_run("Risk score = 65% classifier probability + 35% anomaly score").bold = True
add_bullets(document, [
    "Pressure and flow behavior are especially important because changes can indicate disruption or loss of containment.",
    "Rolling features compare a reading with its recent and longer-term segment behavior.",
    "SHAP values explain which features pushed a prediction higher or lower.",
])
document.add_paragraph("The model now also calculates IQR outlier flags for pressure, flow, temperature, pump speed, and energy consumption. The bounds are learned from the training period using the usual 1.5 × IQR rule and reused during prediction. These flags help identify unusual sensor behavior; they are not proof of a leak and do not directly set the Normal, Warning, or Critical alert thresholds.")

add_heading(document, "8. Dashboard Workflow", 1)
add_numbered(document, [
    "Open the Live Monitor tab.",
    "Review the network status, latest risk, warnings, critical alerts, and data-quality panel.",
    "Use the demo topology or segment list to select a segment.",
    "Review the schematic, gauge, sensor readouts, risk timeline, and contributing factors.",
    "Use the detection log to acknowledge or resolve an alert. The action is saved in SQLite.",
    "Use the telemetry chart to compare pressure and flow over the selected segment history.",
    "Use the What-If Simulator to test a manual reading or scenario preset.",
    "Use Model Performance to review holdout metrics and feature importance.",
])

add_heading(document, "9. What Does a Segment Mean?", 1)
document.add_paragraph("A segment is a monitored unit identified by segment_id. In the current sample there are 50 segments. Each segment has its own sensor history and is scored separately so that rolling statistics describe that segment's behavior.")
add_bullets(document, [
    "SEG-001 through SEG-050 are sample identifiers, not confirmed Titas asset IDs.",
    "The segment list is ranked using historical incidents and risk information.",
    "Selecting a segment replays its readings one by one, similar to a historical live feed.",
    "The current topology is a demo schematic arranged by segment number, not a geographic map.",
    "A real Titas map requires an approved GIS export with coordinates and upstream/downstream relationships.",
])

add_heading(document, "10. Current Dataset and Metrics", 1)
add_table(document, ["Item", "Current value"], [
    ("SCADA readings", "1,000"),
    ("Segments", "50"),
    ("Raw fields", "13"),
    ("Engineered model features after retraining", "72"),
    ("Training rows", "840"),
    ("Chronological test rows", "160"),
    ("Holdout AUC", "0.9926"),
    ("Holdout accuracy", "97.5%"),
    ("Time-series CV F1", "0.892"),
])
document.add_paragraph("These metrics come from the current sample dataset and should not be presented as Titas production accuracy. The original sample had 270 collision groups involving 678 rows with repeated segment timestamps. The processed demo copy now preserves all rows and applies documented one-second offsets within each collision group; this is a synthetic-data repair, not a substitute for a real historian sequence field.")

add_heading(document, "11. How to Run the Project", 1)
document.add_paragraph("From PowerShell, in the project root:")
code = document.add_paragraph()
code.style = "No Spacing"
for line in [
    "py -m pip install -r requirements.txt",
    "py scripts\\repair_sample_timestamps.py",
    "py scripts\\audit_data.py --data data\\processed\\scada_pipeline_repaired.csv",
    "$env:PIPELINE_RISK_MODEL_PATH = \"D:\\ML Project\\Pipeline risk\\artifacts\\models\\pipeline_risk_model.pkl\"",
    "py webapp\\app.py",
]:
    run = code.add_run(line + "\n")
    run.font.name = "Consolas"
    run.font.size = Pt(9)

document.add_paragraph("Open http://127.0.0.1:5000 in a browser. Keep the terminal running while using the dashboard.")
document.add_paragraph("To train a new model, run the sample repair script first, then run py scripts\\train_model.py. For real Titas data, use the historian's true sequence or timestamp instead of the sample repair rule.")

add_heading(document, "12. What Must Change for Real Titas Data?", 1)
add_bullets(document, [
    "Replace sample tag names in config/titas_scada.yml with the approved Titas historian or SCADA tag dictionary.",
    "Confirm pressure, flow, temperature, energy, and speed units.",
    "Provide an approved GIS export for real topology and asset locations.",
    "Define sensor quality flags, maintenance states, and stale-data rules.",
    "Use verified leak, maintenance, sensor-fault, and operational labels.",
    "Review alarm_triggered for target leakage before retraining.",
    "Run in read-only shadow mode before any operational alert integration.",
    "Add access control, audit logging, backups, HTTPS, and protected model/data storage.",
])

add_heading(document, "13. Simple Team Takeaway", 1)
document.add_paragraph("This project is a working prototype of a SCADA risk-monitoring system. It already demonstrates the complete flow from sensor data to machine-learning score to operator dashboard. The next major step is not adding more charts; it is replacing the sample data with a validated Titas data contract, repairing timestamp chronology, connecting approved GIS topology, and validating the model against verified historical incidents.")

note = document.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run("Prepared for project team onboarding")
run.italic = True
run.font.color.rgb = RGBColor(90, 100, 110)

document.save(OUTPUT)
print(OUTPUT)
